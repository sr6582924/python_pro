# -*- coding: utf-8 -*-
"""
-------------------------------------------------
   File Name：     test1
   Description :
   Author :       ming
   date：          2018/12/27
-------------------------------------------------
   Change Activity:
                   2018/12/27:
-------------------------------------------------
"""
import uuid
import docx
import re

from easydb import OpenDB

big_titles = ['一、', '二、', '三、']
answer_element = ['a', 'A', 'b', 'B', 'c', 'C', 'd', 'D']
half_answer = ['b', 'B', 'd', 'D']
half_answer2 = ['a', 'A', 'c', 'C']

def listContains(list, str):
    for title in list:
        if title in str:
            return True
    return False

def checkQuestion(str, index):
    title =  "{0}．".format(index)
    title2 = "{0}.".format(index)
    return title in str or title2 in str

def checkanswer(list, str):
    for title in list:
        if str and title == str.strip()[0]:
            return True
    return False

def insertQuestion(**kwargs):
    with OpenDB() as db:
        sql = "insert into question(id, questionContent, questionDesc, questionType)values('{id}', '{questionContent}', '{questionDesc}', '{questionType}')"
        print(kwargs)
        sql = sql.format(**kwargs)
        print(sql)
        db.execute(sql)

def insertAnswer(**kwargs):
    with OpenDB() as db:
        sql = "insert into answer(id, answer, answerType, questionId, isCorrect, answerCode, sort)values('{id}', '{answer}', '{answerType}', '{questionId}', '{isCorrect}', '{answerCode}', '{sort}')"
        sql = sql.format(**kwargs)
        db.execute(sql)

def isCorect(corrects, answer):
    print("corrects:" + str(corrects) + "answer:" + answer)
    answerNum = answer[0:1]
    return answerNum in corrects

def parseDoc():
    doc = docx.Document("C:\\Users\\ming\\Desktop\\testb\\ma2.docx")
    print("段落数:" + str(len(doc.paragraphs)))  # 段落数为13，每个回车隔离一段
    paragraphs = doc.paragraphs
    questionIndex = 1
    corrects = ''
    questionDesc = ''
    questionType = 0
    sort = 0
    for paragraph in paragraphs:
        title1 = paragraph.text
        if listContains(big_titles, title1):
            print("title" + title1)
            questionDesc = title1
            questionIndex = 1
            questionType += 1
            print(title1)
        if checkQuestion(title1, questionIndex):
            print("question" + title1)
            text = title1
            aa = re.findall("（(.+?)）|\((.+?)\)", text)
            corrects = [i for i in aa[len(aa) -1] if i][0]
            split = re.split("(\(|（)", text)[-1]
            noanswer = text.replace(split, "")[0:-1]
            print(corrects)
            questionId = str(uuid.uuid1())
            insertQuestion(id=questionId, questionContent=noanswer, questionDesc=questionDesc, questionType=questionType)
            questionIndex += 1
            sort = 0
        if checkanswer(answer_element, title1):
            if checkanswer(half_answer2, title1) and listContains(['b．', 'B．','b.', 'B.'], title1):
                answer = re.split("([a-d]|[A-D].|．)", title1)
            elif checkanswer(half_answer2, title1) and listContains([ 'd．', 'D．','d.', 'D.'], title1):
                answer = re.split("([a-d]|[A-D].|．)", title1)
            else:
                answer = title1
            if questionType <= 2:
                if isinstance(answer, list):
                    print("answer:" + str(answer))
                    for ii in range(2, len(answer), 2):
                        ans = answer[ii - 1] + answer[ii]
                        print("ans:" + str(answer))
                        isCorrect = 1 if isCorect(corrects, ans) else 0
                        answerCode = ans[0:1]
                        sort += 1
                        insertAnswer(id=str(uuid.uuid1()), answer=ans.strip(),
                                     answerType=questionType,
                                     questionId=questionId,
                                     isCorrect=isCorrect,
                                     answerCode=answerCode,
                                     sort=sort)
                else:
                    print(answer)
                    isCorrect = 1 if isCorect(corrects, answer) else 0
                    answerCode = answer[0:1]
                    sort += 1
                    insertAnswer(id=str(uuid.uuid1()), answer= answer.strip(),
                                 answerType=questionType,
                                 questionId=questionId,
                                 isCorrect=isCorrect,
                                 answerCode=answerCode,
                                 sort=sort)

if __name__ == '__main__':
    parseDoc()

